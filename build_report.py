"""
Generate ArtinAzma project audit report (RTL Persian Word document).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PERSIAN_FONT = "Vazirmatn"


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_run_font(run, size=11, bold=False, color=None, font=PERSIAN_FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    sz = OxmlElement("w:szCs")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    if bold:
        b = OxmlElement("w:bCs")
        rPr.append(b)


def add_heading(doc, text, level=1):
    sizes = {1: 20, 2: 16, 3: 13}
    colors = {1: (29, 78, 216), 2: (37, 99, 235), 3: (30, 64, 175)}
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    set_run_font(run, size=sizes[level], bold=True, color=colors[level])
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_run_font(r1, size=11, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size=11)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11)
    return p


def add_severity_box(doc, label, color_rgb, text):
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(2.8)
    t.columns[1].width = Cm(13.5)
    cell_label = t.cell(0, 0)
    cell_text = t.cell(0, 1)
    tcPr = cell_label._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_rgb)
    tcPr.append(shd)
    p1 = cell_label.paragraphs[0]
    set_rtl(p1)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(label)
    set_run_font(r1, size=10, bold=True, color=(255, 255, 255))
    p2 = cell_text.paragraphs[0]
    set_rtl(p2)
    r2 = p2.add_run(text)
    set_run_font(r2, size=11, bold=True)
    doc.add_paragraph()


def add_code(doc, text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(180, 30, 60)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        p = cell.paragraphs[0]
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=11, bold=True, color=(255, 255, 255))
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1D4ED8")
        tcPr.append(shd)
    for ridx, row in enumerate(rows, start=1):
        for cidx, val in enumerate(row):
            cell = t.cell(ridx, cidx)
            p = cell.paragraphs[0]
            set_rtl(p)
            r = p.add_run(str(val))
            set_run_font(r, size=10.5)
    doc.add_paragraph()


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = PERSIAN_FONT
    style.font.size = Pt(11)

    # Title
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("گزارش وضعیت و پیشنهاد بهبود پروژه")
    set_run_font(run, size=24, bold=True, color=(29, 78, 216))

    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ArtinAzma Expert Assistant — دستیار هوشمند آرتین")
    set_run_font(run, size=14, color=(75, 85, 99))

    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("تاریخ گزارش: ۲۸ مه ۲۰۲۶")
    set_run_font(run, size=11, color=(107, 114, 128))

    doc.add_paragraph()

    add_heading(doc, "۱. خلاصه مدیریتی", level=1)
    add_para(
        doc,
        "پروژه ArtinAzma Expert Assistant در وضعیت عملیاتی قابل‌قبول و در حال "
        "توسعه فعال است؛ ۲۶۳ commit ثبت‌شده، CI/CD از طریق GitHub Actions راه‌اندازی "
        "شده، Docker و docker-compose آماده deploy است و معماری دو بخشی "
        "(FastAPI + Next.js 16) به‌خوبی تفکیک شده است. با این حال چند آسیب‌پذیری "
        "امنیتی جدی و چندین بدهی فنی (technical debt) قابل توجه شناسایی شد که "
        "قبل از deploy روی سرور Production باید برطرف شود.",
    )
    add_para(
        doc,
        "نمره کلی پروژه: ۷ از ۱۰ — عملکرد و قابلیت‌ها قوی، اما امنیت و ساختار "
        "کد نیازمند پاک‌سازی فوری است.",
        bold=True,
        color=(180, 50, 50),
    )

    add_heading(doc, "۲. تصویر کلی پروژه", level=1)
    add_heading(doc, "۲.۱ اندازه و حجم کد", level=2)
    add_table(
        doc,
        ["بخش", "تعداد فایل / خط", "توضیح"],
        [
            ["Backend (Python)", "۲۸ فایل / ~۱۵٬۱۰۰ خط", "FastAPI، چند سرویس"],
            ["Frontend (TS/TSX)", "~۵۰ فایل / ~۱۶٬۶۰۰ خط", "Next.js 16، React 19"],
            ["تست‌های Backend", "۲ فایل، حدود ۲۰ تست", "pytest"],
            ["Knowledge base", "۴۱۳ مگابایت JSON", "بسیار بزرگ – نیاز به مهاجرت"],
            ["دیتابیس", "SQLite (~۱.۶ مگابایت)", "۱۱ جدول، indexها OK"],
        ],
    )

    add_heading(doc, "۲.۲ نقاط قوت", level=2)
    add_bullet(doc, "معماری تمیز و لایه‌بندی‌شده با تفکیک سرویس‌ها (intent، knowledge، AI، …).")
    add_bullet(doc, "پشتیبانی کامل دو زبان فارسی/انگلیسی و RTL در فرانت‌اند.")
    add_bullet(doc, "Two-tier retrieval (local search + vector embeddings) با fallback هوشمند.")
    add_bullet(doc, "JWT برای احراز هویت مشتری + Rate Limiting هوشمند چندسطحی.")
    add_bullet(doc, "وجود Health Check، Sentry integration، Logging چرخشی و Push Notification.")
    add_bullet(doc, "PWA کامل با service worker، manifest و حالت آفلاین.")
    add_bullet(doc, "CI پایه روی GitHub Actions (lint + typecheck + build).")
    add_bullet(doc, "راه‌حل خلاقانه DoH برای دور زدن بلاکینگ DNS در ایران برای اتصال به OpenAI.")

    add_heading(doc, "۳. مسائل امنیتی (Critical)", level=1)

    add_severity_box(doc, "بحرانی", "DC2626",
                     "کلید ادمین به‌صورت plain-text در کد فرانت‌اند fallback شده است.")
    add_para(doc,
             "در دو فایل زیر، اگر متغیر محیطی ADMIN_API_KEY تنظیم نشده باشد، یک "
             "fallback هاردکدشده بکار می‌رود:")
    add_code(doc, "frontend/src/app/api/admin/customers/route.ts:6")
    add_code(doc, "frontend/src/app/api/admin/customers/[id]/[action]/route.ts:6")
    add_code(doc, 'const ADMIN_API_KEY = process.env.ADMIN_API_KEY || "ArtinAzma1*2#amzanitra!@#";')
    add_para(doc,
             "این کلید در تاریخچه git ذخیره شده و در هر buildای که env درست تنظیم "
             "نشده باشد، به‌عنوان کلید معتبر پذیرفته می‌شود. باید fallback حذف و "
             "در صورت نبود کلید، خطا بازگردانده شود. همچنین کلید فعلی باید rotate شود.",
             bold=True)

    add_severity_box(doc, "بحرانی", "DC2626",
                     "Bypass کامل احراز هویت مشتری در سمت فرانت‌اند.")
    add_para(doc,
             "محافظت از مسیرهای /assistant، /analyze، /customer-dashboard در فایل "
             "frontend/src/proxy.ts فقط بررسی می‌کند که کوکی artin_customer_auth "
             'برابر "logged_in" باشد. ولی این کوکی توسط خود کلاینت در صفحه login ست می‌شود:')
    add_code(doc, 'document.cookie = "artin_customer_auth=logged_in; path=/; max-age=86400";')
    add_para(doc,
             "هر کاربری بدون login می‌تواند با یک خط در DevTools این کوکی را ست "
             "کرده و وارد صفحات محافظت‌شده شود. خوشبختانه backend با JWT جداگانه "
             "محافظت می‌شود ولی روت‌گاردِ فرانت‌اند کاملاً بی‌اثر است. باید این "
             "کوکی به HttpOnly تبدیل شود و فقط توسط API روت سرور بعد از login واقعی ست شود.",
             bold=True)

    add_severity_box(doc, "بحرانی", "DC2626", "Session Token هاردکد شده در fallback.")
    add_para(doc,
             "در proxy.ts و admin-status/route.ts اگر ADMIN_SESSION_TOKEN تنظیم "
             'نشده باشد، مقدار ثابت "artin-local-admin-session" پذیرفته می‌شود. '
             "این یعنی هر کسی که این مقدار را در کوکی ست کند، دسترسی ادمین به دست می‌آورد.")

    add_severity_box(doc, "بحرانی", "DC2626", "فایل‌های حساس در workspace قرار دارند.")
    add_para(doc,
             "فایل backend/.env و backend/google-service-account.json در workspace "
             "موجود هستند. هرچند در .gitignore اضافه شده‌اند، باید مطمئن شد در "
             "تاریخچه git قبلاً commit نشده‌اند. اجرا کنید:")
    add_code(doc, "git log --all --diff-filter=A -- backend/.env backend/google-service-account.json")
    add_para(doc, "اگر پاسخی برگشت، کلیدها باید بلافاصله rotate شوند.")

    add_severity_box(doc, "بالا", "EA580C", "غیرفعال‌سازی verify SSL در DNS over HTTPS.")
    add_para(doc,
             "در backend/ai_service.py (خطوط ۶۳-۶۶) برای DoH اتصال به Cloudflare/Google، "
             "بررسی hostname و گواهی SSL غیرفعال شده. این به‌دلیل اتصال با IP خام "
             "منطقی است، اما حمله MITM روی همان درخواست‌های DoH ممکن می‌شود. "
             "می‌توان گواهی Cloudflare را pinned کرد یا فقط hostname verification "
             "را غیرفعال کرد اما verify_mode را CERT_REQUIRED گذاشت با CA bundle مناسب.")

    add_severity_box(doc, "بالا", "EA580C", "ADMIN_PASSWORD به‌صورت plain text در env.")
    add_para(doc,
             "در api/admin-login/route.ts مقایسه password با مقدار env به‌صورت "
             "plain string انجام می‌شود. باید hash bcrypt در env ذخیره شود و در "
             "سمت سرور verify شود، یا حداقل از مقایسه‌ی timing-safe (مثل "
             "crypto.timingSafeEqual) استفاده شود.")

    add_heading(doc, "۴. مسائل ساختار و کیفیت کد", level=1)

    add_severity_box(doc, "بالا", "F59E0B", "فایل main.py با ۲۹۰۴ خط — God Object.")
    add_para(doc,
             "بیش از ۸۰ endpoint در یک فایل قرار دارد. باید با APIRouter به "
             "ماژول‌های زیر تقسیم شود:")
    add_bullet(doc, "routes/chat.py — تمام endpointهای /chat و /chat/stream")
    add_bullet(doc, "routes/knowledge.py — مدیریت knowledge base و Google Drive")
    add_bullet(doc, "routes/analysis.py — تحلیل فایل و عکس")
    add_bullet(doc, "routes/health.py — health/system/status")
    add_bullet(doc, "schemas/ — تمام Pydantic models")

    add_severity_box(doc, "بالا", "F59E0B",
                     "فایل db_service.py با ۱۸۸۹ خط — repositoryهای جدا لازم است.")
    add_para(doc,
             "همه عملیات CRUD برای ۱۱ جدول در یک فایل قرار دارد. تقسیم به "
             "repositories/customer_repo.py، repositories/chat_repo.py و … "
             "خوانایی و قابلیت تست را به‌مراتب بهبود می‌دهد. اضافه‌کردن یک ORM "
             "سبک مثل SQLAlchemy 2.0 یا SQLModel نیز ارزش بررسی دارد.")

    add_severity_box(doc, "بالا", "F59E0B",
                     "فایل knowledge_vectors.json با حجم ۴۱۳ مگابایت.")
    add_para(doc,
             "تمام embedding ها در یک فایل JSON نگهداری می‌شوند که در حافظه load "
             "می‌شود. این به مرور باعث:")
    add_bullet(doc, "افزایش زمان startup سرور به چند ده ثانیه")
    add_bullet(doc, "مصرف بالای RAM (احتمالاً > 1.5 GB)")
    add_bullet(doc, "خطر فساد فایل در نوشتن همزمان (race condition)")
    add_bullet(doc, "ناتوانی در scale-out می‌شود")
    add_para(doc,
             "خوشبختانه qdrant_service.py و migrate_to_qdrant.py از قبل آماده "
             "شده‌اند. مهاجرت کامل به Qdrant باید اولویت دار باشد و فایل JSON فقط "
             "به‌عنوان fallback یا backup استفاده شود.", bold=True)

    add_severity_box(doc, "متوسط", "0EA5E9", "کش دوگانه پاسخ‌ها — duplicated code.")
    add_para(doc,
             "یک _ResponseCache در main.py و یک _RESPONSE_CACHE در ai_service.py "
             "وجود دارد که عمل مشابهی انجام می‌دهند. باید در یک ماژول جداگانه "
             "مثل services/cache.py یکی شوند.")

    add_severity_box(doc, "متوسط", "0EA5E9", "سه فایل requirements متناقض.")
    add_para(doc,
             "requirements.txt، requirements_clean.txt و requirements_fixed.txt "
             "هر سه در پوشه backend وجود دارند. باید فقط یکی نگه داشته شود و "
             "بقیه حذف شوند. همچنین چند نسخه‌ی غیرواقع‌بینانه دیده می‌شود "
             "(مثلاً pandas==3.0.2، openai==2.33.0، fastapi==0.136.1) که باید "
             "با pip-tools/uv lock کنترل شوند.")

    add_severity_box(doc, "متوسط", "0EA5E9", "اسکریپت fix_artin_ui.py در ریشه پروژه.")
    add_para(doc,
             "فایل ۱۸ کیلوبایتی fix_artin_ui.py در ریشه قرار دارد و به نظر می‌رسد "
             "یک اسکریپت یک‌بارمصرف باشد. باید به scripts/ منتقل یا حذف شود.")

    add_severity_box(doc, "متوسط", "0EA5E9",
                     "تست‌ها از main.py به‌صورت سراسری import می‌کنند.")
    add_para(doc,
             "از from main import is_specific_product_or_model_question استفاده شده "
             "است؛ این یعنی هر شکست در lifespan هنگام تست، باعث شکست همه تست‌ها "
             "می‌شود. بهتر است این توابع به ماژول utils/text_classifier.py منتقل شوند.")

    add_heading(doc, "۵. مسائل عملکرد (Performance)", level=1)
    add_bullet(doc,
               "Vector store JSON باعث latency بالا در جستجوی محلی می‌شود — "
               "مهاجرت به Qdrant ضروری است.")
    add_bullet(doc,
               "SQLite برای محیط چندکاربره مناسب نیست؛ در صورت رشد کاربران به "
               "PostgreSQL مهاجرت کنید (تغییر کوچک با SQLAlchemy).")
    add_bullet(doc,
               "page.tsx در assistant با ۱۹۵۳ خط، کامپوننت‌های فرعی باید جدا شوند "
               "تا re-render مدیریت بهتری داشته باشد.")
    add_bullet(doc,
               "تابع save_expert_question در chat هر بار سینکرون اجرا می‌شود؛ "
               "بهتر است به background task FastAPI منتقل شود تا زمان پاسخ کاهش یابد.")
    add_bullet(doc,
               "_request_logger فعلی روی هر درخواست trigger می‌شود؛ روی مسیرهای "
               "پراستفاده (مثل /health) باید skip شود.")
    add_bullet(doc,
               "Service worker قدیمی pre-cache می‌کند ولی استراتژی stale-while-revalidate "
               "برای صفحات داینامیک تعریف نشده — کاربران بعد از deploy ممکن است "
               "نسخه قدیمی ببینند.")

    add_heading(doc, "۶. UX و فرانت‌اند", level=1)
    add_bullet(doc,
               "JWT در localStorage ذخیره می‌شود (آسیب‌پذیر به XSS) — انتقال به "
               "HttpOnly cookie توصیه می‌شود.")
    add_bullet(doc,
               "/customer-login و /customer-register خودشان کوکی auth را ست می‌کنند — "
               "این منطق باید فقط در API route سرور باشد.")
    add_bullet(doc,
               "اضافه‌کردن React Query یا SWR برای cache و refetch هوشمند به‌جای "
               "fetch دستی توصیه می‌شود.")
    add_bullet(doc,
               "هیچ Error Boundary در سطح page وجود ندارد (فقط ErrorBoundary "
               "سراسری) — افتادن یک کامپوننت تمام صفحه را خراب می‌کند.")
    add_bullet(doc,
               "فایل assistant-polish.css در کنار globals.css سردرگمی ایجاد می‌کند — "
               "یکپارچه‌سازی توصیه می‌شود.")
    add_bullet(doc,
               "تست PWA روی iOS Safari توصیه می‌شود (manifest و sw.js).")

    add_heading(doc, "۷. DevOps و Deployment", level=1)
    add_bullet(doc, "CI فعلی فقط lint و build اجرا می‌کند — اجرای pytest نیز اضافه شود.")
    add_bullet(doc,
               "Dockerfile بک‌اند بدون non-root user اجرا می‌شود — اضافه‌کردن "
               "USER appuser توصیه می‌شود.")
    add_bullet(doc, "هیچ healthcheck برای frontend در docker-compose تعریف نشده.")
    add_bullet(doc,
               "فایل‌های sitemap (page-sitemap.xml و …) در پوشه backend رها شده‌اند — "
               "باید به nginx یا public/ منتقل شوند.")
    add_bullet(doc,
               "بدون reverse proxy (nginx/caddy) در docker-compose — TLS termination "
               "تعریف نشده.")
    add_bullet(doc,
               "ADMIN_PASSWORD و سایر secrets از env فایل خوانده می‌شود؛ توصیه "
               "می‌شود از Docker secrets یا vault استفاده شود.")

    add_heading(doc, "۸. تست و کیفیت", level=1)
    add_bullet(doc, "پوشش تست بک‌اند بسیار پایین (~۲۰ تست برای ~۱۵٬۰۰۰ خط کد).")
    add_bullet(doc, "هیچ تستی برای frontend وجود ندارد — Vitest یا Playwright اضافه شود.")
    add_bullet(doc,
               "سرویس‌های پیچیده مثل ai_service، intent_service، knowledge_service "
               "بدون unit test هستند.")
    add_bullet(doc,
               "حذف کامل بدنه‌های try/except که فقط print می‌کنند؛ باید "
               "logger.exception شود تا Sentry آن‌ها را بگیرد.")
    add_bullet(doc, "اضافه کردن mypy/pyright برای type checking در CI.")

    add_heading(doc, "۹. نقشه راه پیشنهادی (اولویت‌بندی‌شده)", level=1)

    add_heading(doc, "هفته ۱ — رفع آسیب‌پذیری‌های بحرانی", level=2)
    add_bullet(doc, "حذف fallback کلید ادمین در api/admin/customers/* و rotate کردن کلید فعلی.", bold_prefix="①")
    add_bullet(doc, "تبدیل artin_customer_auth به کوکی HttpOnly که فقط API server ست می‌کند.", bold_prefix="②")
    add_bullet(doc, "حذف fallback Session Token در proxy.ts و admin-status.", bold_prefix="③")
    add_bullet(doc, "بررسی تاریخچه git برای فایل .env و service account JSON. در صورت یافت‌شدن، rotate تمام کلیدها.", bold_prefix="④")
    add_bullet(doc, "تبدیل ADMIN_PASSWORD به bcrypt hash.", bold_prefix="⑤")

    add_heading(doc, "هفته ۲-۳ — مهاجرت Qdrant و پاک‌سازی", level=2)
    add_bullet(doc, "اجرای migrate_to_qdrant.py در staging و سپس production.", bold_prefix="①")
    add_bullet(doc, "حذف فایل ۴۱۳MB knowledge_vectors.json از mainline.", bold_prefix="②")
    add_bullet(doc, "ادغام سه فایل requirements و قفل با uv یا pip-tools.", bold_prefix="③")
    add_bullet(doc, "حذف fix_artin_ui.py و sitemap.xmlها از پوشه backend.", bold_prefix="④")

    add_heading(doc, "هفته ۴-۶ — Refactor و افزایش تست", level=2)
    add_bullet(doc, "شکستن main.py به ۴-۵ روتر مجزا.", bold_prefix="①")
    add_bullet(doc, "شکستن db_service.py به repository های جدا، اضافه‌کردن SQLAlchemy.", bold_prefix="②")
    add_bullet(doc, "اضافه کردن pytest در CI و رساندن coverage به > ۵۰٪.", bold_prefix="③")
    add_bullet(doc, "شکستن assistant/page.tsx به کامپوننت‌های کوچک‌تر.", bold_prefix="④")

    add_heading(doc, "ماه ۲ — Scale و قابلیت اطمینان", level=2)
    add_bullet(doc, "مهاجرت SQLite → PostgreSQL با Alembic برای migrationها.", bold_prefix="①")
    add_bullet(doc, "اضافه‌کردن Redis برای cache مشترک بین instanceها.", bold_prefix="②")
    add_bullet(doc, "اضافه‌کردن nginx یا Caddy جلوی هر دو سرویس + TLS.", bold_prefix="③")
    add_bullet(doc, "اضافه‌کردن Prometheus metrics + Grafana dashboard.", bold_prefix="④")

    add_heading(doc, "۱۰. خلاصه و جمع‌بندی نهایی", level=1)
    add_para(doc,
             "این پروژه از نظر قابلیت‌ها و معماری کلی در سطح حرفه‌ای است و قابلیت "
             "تبدیل شدن به یک محصول SaaS کامل را دارد. نقاط قوت اصلی شامل تجربه "
             "کاربری تخصصی، پشتیبانی کامل از زبان فارسی، retrieval چندلایه و ادغام "
             "موفق با OpenAI (حتی با محدودیت‌های ایران) است.")
    add_para(doc,
             "نقاط ضعف به‌طور عمده در حوزه‌ی امنیت اولیه و بدهی فنی هستند. اگر این "
             "پروژه قرار است در محیط تولید با کاربران واقعی منتشر شود، حداقل موارد "
             "هفته ۱ باید قبل از launch انجام شود — در غیر این صورت ریسک نشت کلید "
             "ادمین و bypass دسترسی بسیار بالا است.",
             bold=True)
    add_para(doc,
             "با اجرای نقشه راه فوق ظرف ۲-۳ ماه، پروژه می‌تواند به سطح "
             "production-ready کامل برسد.",
             color=(29, 78, 216), bold=True)

    doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— پایان گزارش —")
    set_run_font(run, size=10, color=(150, 150, 150))

    out_path = "/sessions/gracious-wizardly-goodall/mnt/artinazma-expert-assistant/گزارش-وضعیت-پروژه.docx"
    doc.save(out_path)
    print(f"saved {out_path}")


if __name__ == "__main__":
    main()
