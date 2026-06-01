"""
Generate ArtinAzma project status DIFF report (before / after).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PERSIAN_FONT = "Vazirmatn"


def set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_run_font(run, size=11, bold=False, color=None, font=PERSIAN_FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    sz = OxmlElement("w:szCs")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)


def heading(doc, text, level=1):
    sizes = {1: 20, 2: 16, 3: 13}
    colors = {1: (29, 78, 216), 2: (37, 99, 235), 3: (30, 64, 175)}
    p = doc.add_paragraph()
    set_rtl(p)
    r = p.add_run(text)
    set_run_font(r, size=sizes[level], bold=True, color=colors[level])
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)


def para(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_run_font(r1, size=11, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size=11)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11)


def status_row(doc, label_color, label, text):
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(2.5)
    t.columns[1].width = Cm(14)
    cl = t.cell(0, 0)
    ct = t.cell(0, 1)
    tcPr = cl._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), label_color)
    tcPr.append(shd)
    p1 = cl.paragraphs[0]
    set_rtl(p1)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(label)
    set_run_font(r1, size=10, bold=True, color=(255, 255, 255))
    p2 = ct.paragraphs[0]
    set_rtl(p2)
    r2 = p2.add_run(text)
    set_run_font(r2, size=11, bold=True)
    doc.add_paragraph()


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        p = cell.paragraphs[0]
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=11, bold=True, color=(255, 255, 255))
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1D4ED8")
        tcPr.append(shd)
    for ridx, row in enumerate(rows, start=1):
        for cidx, val in enumerate(row):
            cell = t.cell(ridx, cidx)
            p = cell.paragraphs[0]
            set_rtl(p)
            r = p.add_run(str(val))
            set_run_font(r, size=10.5)
    doc.add_paragraph()


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)
    doc.styles["Normal"].font.name = PERSIAN_FONT
    doc.styles["Normal"].font.size = Pt(11)

    # Title
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("گزارش وضعیت — بعد از اعمال اصلاحات")
    set_run_font(r, size=22, bold=True, color=(29, 78, 216))

    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ArtinAzma Expert Assistant — مقایسه قبل و بعد")
    set_run_font(r, size=14, color=(75, 85, 99))

    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("تاریخ گزارش: ۲۹ مه ۲۰۲۶")
    set_run_font(r, size=11, color=(107, 114, 128))

    doc.add_paragraph()

    # ── Executive summary ───────────────────────────────────
    heading(doc, "۱. خلاصه — کار خیلی خوبی شده!", level=1)
    para(doc,
         "از گزارش قبلی تا الان حجم بسیار قابل‌توجهی از اصلاحات انجام شده. تقریباً "
         "تمام مسائل بحرانی امنیتی برطرف شده‌اند، main.py و db_service.py کاملاً "
         "refactor شده‌اند، Alembic، PostgreSQL، Redis، Qdrant، nginx، Prometheus "
         "و Grafana به docker-compose اضافه شده و تست‌ها سه برابر شده‌اند.")
    para(doc, "نمره فعلی: ۸.۵ از ۱۰ (قبل: ۷ از ۱۰)", bold=True, color=(34, 139, 34))

    # ── What was fixed ──────────────────────────────────────
    heading(doc, "۲. مواردی که رفع شده (✓)", level=1)

    heading(doc, "۲.۱ امنیت (تمام مسائل بحرانی رفع شد)", level=2)
    status_row(doc, "16A34A", "رفع شد",
               "کلید ادمین hardcoded حذف شد — حالا اگر env نباشد throw error می‌کند.")
    status_row(doc, "16A34A", "رفع شد",
               "Customer auth bypass رفع شد — کوکی HMAC signed httpOnly در /api/customer-session.")
    status_row(doc, "16A34A", "رفع شد",
               "Session token fallback حذف شد — proxy.ts fail-closed عمل می‌کند.")
    status_row(doc, "16A34A", "رفع شد",
               "ADMIN_PASSWORD حالا با timingSafeEqual + SHA-256 مقایسه می‌شود.")
    status_row(doc, "16A34A", "رفع شد",
               "کوکی artin_customer_auth حالا httpOnly است و فقط API server ست می‌کند.")

    heading(doc, "۲.۲ ساختار کد (refactor عالی)", level=2)
    table(doc,
          ["فایل", "قبل", "بعد", "نتیجه"],
          [
              ["main.py", "۲۹۰۴ خط", "۲۳۷ خط", "✓ شکسته شد به ۷ روتر"],
              ["db_service.py", "۱۸۸۹ خط", "۴۵۶ خط", "✓ به ۸ repository تقسیم شد"],
              ["routes/", "۲ فایل", "۹ فایل", "✓ chat, knowledge, analysis, …"],
              ["repositories/", "وجود نداشت", "۸ ماژول", "✓ pattern تمیز"],
              ["schemas/", "همه در main.py", "models.py جدا", "✓ Pydantic models جدا"],
              ["utils/", "وجود نداشت", "chat_utils + deps", "✓ helperهای مشترک"],
          ])

    heading(doc, "۲.۳ DevOps (تحول کامل)", level=2)
    bullet(doc, "PostgreSQL با healthcheck اضافه شد.")
    bullet(doc, "Redis با maxmemory policy اضافه شد.")
    bullet(doc, "Qdrant container اضافه شد (vector DB واقعی).")
    bullet(doc, "nginx reverse proxy + TLS termination اضافه شد.")
    bullet(doc, "Certbot برای Let's Encrypt اضافه شد.")
    bullet(doc, "Prometheus + Grafana برای monitoring اضافه شد.")
    bullet(doc, "Alembic برای migration دیتابیس راه‌اندازی شد (با initial schema).")
    bullet(doc, "Prometheus instrumentation در FastAPI اضافه شد (/metrics).")
    bullet(doc, "metrics_service.py برای gauge های سفارشی نوشته شد.")
    bullet(doc, "CI حالا pytest را هم اجرا می‌کند (قبلاً فقط lint+build بود).")
    bullet(doc, "CI یک job اضافی برای docker build verification دارد.")
    bullet(doc, "همه سرویس‌ها healthcheck دارند.")
    bullet(doc, "فایل‌های sitemap.xml از backend حذف شدند.")
    bullet(doc, "اسکریپت fix_artin_ui.py حذف شد.")
    bullet(doc, "سه فایل requirements ادغام شدند به یک requirements.txt.")

    heading(doc, "۲.۴ تست‌ها (افزایش قابل توجه)", level=2)
    table(doc,
          ["معیار", "قبل", "بعد"],
          [
              ["تعداد فایل تست", "۲", "۵"],
              ["خطوط تست", "~۲۸۰", "~۹۵۸"],
              ["پوشش", "API basic + services", "+ repositories + chat sessions + knowledge"],
          ])

    # ── Remaining issues ─────────────────────────────────────
    heading(doc, "۳. آنچه هنوز باقی مانده", level=1)

    status_row(doc, "DC2626", "بحرانی",
               "فایل knowledge_vectors.json هنوز ۴۰۵ مگابایت — همراه با backup .bak ۸۱۰ مگابایت")
    para(doc,
         "حتی اگر Qdrant را به docker-compose اضافه کرده‌اید، فایل JSON هنوز سر "
         "جای خودش هست و احتمالاً هنوز توسط knowledge_service یا local_search "
         "خوانده می‌شود. باید migrate_to_qdrant.py اجرا و سپس فایل JSON حذف شود.")
    bullet(doc, "اجرا: python migrate_to_qdrant.py", bold_prefix="→")
    bullet(doc, "بررسی: backend/knowledge_service.py باید فقط از qdrant_service استفاده کند.", bold_prefix="→")
    bullet(doc, "حذف فایل JSON بعد از تأیید موفقیت migration.", bold_prefix="→")
    bullet(doc, "همچنین knowledge_vectors.json.bak اضافی — حذف شود.", bold_prefix="→")

    status_row(doc, "EA580C", "بالا", "Dockerfile بک‌اند هنوز بدون non-root user است")
    para(doc,
         "frontend/Dockerfile به‌درستی USER nextjs دارد ولی backend/Dockerfile "
         "هیچ USER appuser ندارد. در صورت RCE، حمله‌کننده دسترسی root در "
         "container می‌گیرد. سه خط ساده:")
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("RUN adduser --system --uid 1001 appuser\nCOPY --chown=appuser:appuser . .\nUSER appuser")
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(180, 30, 60)

    status_row(doc, "EA580C", "بالا", "غیرفعال‌سازی verify SSL در DoH")
    para(doc,
         "در ai_service.py هنوز ctx.verify_mode = ssl.CERT_NONE برای DoH وجود "
         "دارد. ریسک MITM روی همان درخواست‌های DoH باقی است. می‌توان CA bundle "
         "مناسب گذاشت و فقط hostname verification را disable کرد.")

    status_row(doc, "F59E0B", "متوسط",
               "فایل assistant/page.tsx همچنان ۱۶۰۵ خط")
    para(doc,
         "از ۱۹۵۳ به ۱۶۰۵ خط کاهش یافته ولی همچنان بسیار بزرگ است. باید به "
         "کامپوننت‌های جدا تقسیم شود (ChatMessages، InputBar، ToolPanel، …).")

    status_row(doc, "F59E0B", "متوسط",
               "هیچ تستی برای frontend وجود ندارد")
    para(doc,
         "Vitest + React Testing Library یا Playwright اضافه شود. حداقل برای "
         "اجزای کلیدی مثل MessageBubble، فرم‌های login و chat input.")

    status_row(doc, "F59E0B", "متوسط",
               "دو فایل روتر تکراری در backend/routes/")
    para(doc,
         "هم admin.py + admin_routes.py وجود دارد و هم customers.py + "
         "customer_routes.py. به نظر می‌رسد در حین refactor فایل‌های قدیمی پاک "
         "نشده‌اند. باید بررسی کنید کدام واقعاً استفاده می‌شود و دیگری حذف شود "
         "(main.py فقط customers.py و admin.py را include می‌کند، پس "
         "admin_routes.py و customer_routes.py احتمالاً مرده هستند).")

    status_row(doc, "0EA5E9", "پایین",
               "DATABASE_URL در docker-compose ست شده ولی پروژه هنوز SQLite می‌خواند")
    para(doc,
         "docker-compose.yml متغیر DATABASE_URL را به PostgreSQL ست می‌کند ولی "
         "db_service.py همچنان به sqlite3.connect(\"storage/app.db\") اشاره می‌کند. "
         "Alembic داره راه می‌افته ولی mainline backend هنوز روی SQLite است. باید "
         "db_service به SQLAlchemy/psycopg2 منتقل شود یا حداقل DATABASE_URL را "
         "بخواند.")

    status_row(doc, "0EA5E9", "پایین",
               "JWT همچنان در localStorage ذخیره می‌شود")
    para(doc,
         "هرچند کوکی httpOnly برای session اضافه شده، اما توکن JWT اصلی برای "
         "API backend هنوز در localStorage است (در api.ts). آسیب‌پذیر به XSS. "
         "می‌توان آن را هم به httpOnly cookie منتقل کرد.")

    status_row(doc, "0EA5E9", "پایین",
               "ADMIN_PASSWORD همچنان plain در env")
    para(doc,
         "SHA-256 + timingSafeEqual خوب است ولی password هنوز plain در .env "
         "ذخیره می‌شود. برای production بهتر است bcrypt hash در .env باشد و "
         "verify در سرور.")

    status_row(doc, "0EA5E9", "پایین",
               "فایل‌های حساس در workspace")
    para(doc,
         "backend/.env و backend/google-service-account.json هنوز در پوشه "
         "هستند. اگر تاریخچه git قبلاً آنها را commit کرده باشد، باید با "
         "git filter-repo یا BFG پاک شود.")

    # ── New observations ─────────────────────────────────────
    heading(doc, "۴. نکات جدید — مواردی که در گزارش قبل نبود", level=1)

    bullet(doc,
           "فایل knowledge_vectors.json.bak (۴۰۵ مگابایت) وجود دارد — احتمالاً قبل "
           "از یک refactor backup گرفته شده. باید حذف شود.")
    bullet(doc,
           "alembic.ini هنوز sqlalchemy.url = driver://user:pass@localhost/dbname "
           "دارد. باید با env var از DATABASE_URL خوانده شود.")
    bullet(doc,
           "POSTGRES_PASSWORD در docker-compose default value \"changeme\" دارد — "
           "این در production خطرناک است؛ بهتر است اگر env ست نشده، container fail کند.")
    bullet(doc,
           "Prometheus endpoint /metrics بدون auth در دسترس است — هرچند ports "
           "expose نمی‌شوند، اما در صورت اشتباه در config nginx می‌تواند نشت کند.")
    bullet(doc,
           "Grafana با password \"artinazma_grafana\" به‌عنوان default — باید برای "
           "production تغییر کند.")

    # ── Recommendations ──────────────────────────────────────
    heading(doc, "۵. توصیه اولویت‌بندی‌شده مرحله بعد", level=1)

    heading(doc, "این هفته:", level=2)
    bullet(doc, "اجرای migrate_to_qdrant.py و حذف knowledge_vectors.json + .bak.", bold_prefix="①")
    bullet(doc, "اضافه‌کردن USER appuser به backend/Dockerfile.", bold_prefix="②")
    bullet(doc, "حذف admin_routes.py و customer_routes.py مرده.", bold_prefix="③")

    heading(doc, "هفته بعد:", level=2)
    bullet(doc, "مهاجرت db_service.py از sqlite3 خام به SQLAlchemy + DATABASE_URL.", bold_prefix="①")
    bullet(doc, "اصلاح alembic.ini برای استفاده از env var.", bold_prefix="②")
    bullet(doc, "شکستن assistant/page.tsx به ۴-۵ کامپوننت کوچک‌تر.", bold_prefix="③")

    heading(doc, "ماه بعد:", level=2)
    bullet(doc, "اضافه‌کردن Vitest و چند تست برای کامپوننت‌های کلیدی frontend.", bold_prefix="①")
    bullet(doc, "تبدیل JWT customer به httpOnly cookie.", bold_prefix="②")
    bullet(doc, "حل verify SSL در DoH با CA bundle مناسب.", bold_prefix="③")
    bullet(doc, "تبدیل ADMIN_PASSWORD به bcrypt hash در env.", bold_prefix="④")

    # ── Final summary ────────────────────────────────────────
    heading(doc, "۶. جمع‌بندی نهایی", level=1)
    para(doc,
         "تبریک! این یک کار refactor فوق‌العاده بوده. پروژه از یک کد‌بیس "
         "monolithic با چند آسیب‌پذیری بحرانی، به یک معماری تمیز و حرفه‌ای "
         "تبدیل شده که آماده production است.",
         color=(34, 139, 34), bold=True)
    para(doc,
         "تنها یک کار بحرانی باقی مانده: حذف فایل ۸۱۰ مگابایتی knowledge_vectors. "
         "این بزرگ‌ترین مانع برای deploy روی سرور با RAM محدود است. بقیه موارد "
         "(non-root user، فایل‌های تکراری، تست‌های frontend) همگی کارهای کوچک "
         "روزانه هستند که در یک هفته می‌توان تمام کرد.")
    para(doc,
         "اگر فقط مهاجرت Qdrant را تمام کنید، می‌توانید با اطمینان launch کنید.",
         bold=True, color=(29, 78, 216))

    doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— پایان گزارش —")
    set_run_font(r, size=10, color=(150, 150, 150))

    out = "/sessions/gracious-wizardly-goodall/mnt/artinazma-expert-assistant/گزارش-وضعیت-بعد-از-اصلاحات.docx"
    doc.save(out)
    print(f"saved {out}")


if __name__ == "__main__":
    main()
